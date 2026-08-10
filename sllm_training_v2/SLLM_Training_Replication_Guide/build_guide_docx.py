# -*- coding: utf-8 -*-
"""
Renders guide_source.md -> SLLM_Training_Pipeline_Guide.docx.

This is a copy of report_assets/build_docx.py with one addition: a
line-merging pre-pass that joins soft-wrapped continuation lines back into
a single logical line before parsing, so paragraphs/bullets written across
multiple wrapped lines in the .md source render as one paragraph instead of
splitting into several. Everything else (headings, tables, code blocks,
bullets, inline ** / ` formatting) matches the original renderer exactly,
so any future .docx report in this project keeps a consistent look.
"""
import re
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = sys.argv[1] if len(sys.argv) > 1 else os.path.join(BASE_DIR, "guide_source.md")
OUT_PATH = sys.argv[2] if len(sys.argv) > 2 else os.path.join(BASE_DIR, "SLLM_Training_Pipeline_Guide.docx")

INK = RGBColor(0x0B, 0x0B, 0x0B)
INK_SECONDARY = RGBColor(0x52, 0x51, 0x4E)
ACCENT = RGBColor(0x2A, 0x78, 0xD6)
MUTED = RGBColor(0x89, 0x87, 0x81)

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")

BLOCK_START_RE = re.compile(r"^(#{1,3}\s|```|\|\s?|-\s|\d+\.\s|!\[|---\s*$|---$)")


def merge_wrapped_lines(raw_lines):
    """Join consecutive non-blank lines into one logical line, except where
    a new line starts a new block (heading/code/table/bullet/image/rule) or
    a code fence is open (preserve verbatim)."""
    merged = []
    in_code = False
    buf = None
    for line in raw_lines:
        if line.strip().startswith("```"):
            if buf is not None:
                merged.append(buf)
                buf = None
            merged.append(line)
            in_code = not in_code
            continue
        if in_code:
            merged.append(line)
            continue
        if not line.strip():
            if buf is not None:
                merged.append(buf)
                buf = None
            merged.append(line)
            continue
        if BLOCK_START_RE.match(line.strip()) or buf is None:
            if buf is not None:
                merged.append(buf)
            buf = line.rstrip()
        else:
            buf = buf + " " + line.strip()
    if buf is not None:
        merged.append(buf)
    return merged


def add_inline_runs(paragraph, text, base_size=10.5):
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = paragraph.add_run(chunk[2:-2])
            r.bold = True
            r.font.size = Pt(base_size)
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = paragraph.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(base_size - 0.5)
            r.font.color.rgb = RGBColor(0x8A, 0x2B, 0x06)
        else:
            r = paragraph.add_run(chunk)
            r.font.size = Pt(base_size)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def style_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "D9D8D2")
        borders.append(el)
    tbl_pr.append(borders)


def add_markdown_table(doc, rows):
    n_cols = len(rows[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_table_borders(table)
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            if j >= n_cols:
                break
            cells[j].width = Inches(6.2 / n_cols)
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, val, base_size=9.5)
            if i == 0:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cells[j], "2A78D6")
            else:
                set_cell_shading(cells[j], "FCFCFB" if i % 2 else "F2F2EF")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_heading(doc, text, level):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(14 if level > 1 else 4)
    h.paragraph_format.space_after = Pt(8)
    add_inline_runs(h, text, base_size={1: 20, 2: 15, 3: 12.5}.get(level, 11))
    for r in h.runs:
        r.font.color.rgb = INK if level <= 2 else ACCENT
        r.bold = True
    return h


def add_paragraph(doc, text, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_inline_runs(p, text, base_size=size)
    if italic:
        for r in p.runs:
            r.italic = True
    if color:
        for r in p.runs:
            r.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    add_inline_runs(p, text, base_size=10.5)


def add_image(doc, rel_path, caption=None):
    abs_path = os.path.join(BASE_DIR, rel_path)
    if os.path.isfile(abs_path):
        doc.add_picture(abs_path, width=Inches(6.3))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last.paragraph_format.space_after = Pt(10)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2EF")
    pPr.append(shd)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = INK_SECONDARY


def parse_and_build():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        raw_lines = f.read().split("\n")

    lines = merge_wrapped_lines(raw_lines)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if line.startswith("```"):
            i += 1
            block = []
            while i < n and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, block)
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("!["):
            m = re.match(r"!\[.*?\]\((.+?)\)", line.strip())
            if m:
                add_image(doc, m.group(1))
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for idx, tl in enumerate(table_lines):
                if idx == 1 and re.match(r"^\|?[\s:|-]+\|?$", tl):
                    continue
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            if rows:
                add_markdown_table(doc, rows)
            continue

        if re.match(r"^\d+\.\s+", line.strip()) or line.strip().startswith("- "):
            text = line.strip()
            text = re.sub(r"^\d+\.\s+", "", text)
            text = text[2:] if text.startswith("- ") else text
            add_bullet(doc, text)
            i += 1
            continue

        add_paragraph(doc, line.strip())
        i += 1

    doc.save(OUT_PATH)
    print("Saved", OUT_PATH)


if __name__ == "__main__":
    parse_and_build()
